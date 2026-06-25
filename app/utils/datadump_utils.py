import os
import re
import fitz
from flask import current_app
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document

def create_prefilled_doc(rec):

    folder = current_app.config["DATADUMP_GENERATED_FOLDER"]
    os.makedirs(folder, exist_ok=True)

    doc = Document()
    doc.add_heading("Format for Sharing e-Procurement Data to AG Audit",level=1)

    doc.add_heading("A) Sharing of Data to State Co-ordinator, NIC: ", level=2)

    table = doc.add_table(rows=12, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Remarks"


    fields = [
        ("Name of the Nodal Department/State:", rec.state),
        ("Nodal Department", rec.nodal_dept),
        ("Coordinator", rec.coordinator),
        ("Name of the file shared", rec.file_name),
        ("MD5 Hash of the File Shared", rec.md5_hash),
        ("Size of the File", rec.file_size),
        ("Period Up-to which the data was shared", rec.period_shared),
        ("Postgresql Version in which the Database backup was taken", rec.postgres_version),
        ("Command to retrieve the data in a new Database (DB Name: gepnag, Owner: gepuser)", rec.command_to_restore),
        ("Probable size of database after import", rec.db_size),
        (
            "Sharing of Data to State NIC Coordinator\n"
            "(i) Mode\n"
            "(ii) Date\n"
            "(iii) Name of the Coordinator",
            f"{rec.share_mode or ''}, {str(rec.share_date or '')}, {getattr(rec, 'coordinator_name', '') or ''}"
        ),
    ]

    for row, (label, value) in zip(table.rows[1:], fields):
        row.cells[0].text = label
        row.cells[1].text = str(value or "")



    # ===========
    # Table 2
    # ===========
    doc.add_paragraph("\n")
    doc.add_heading("B). Sharing of eProcurement Data to  State Nodal Agency/Department by State NIC Coordinator", level=2)

    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"


    table2_fields = [
        (
            "Details of Data Dept. Receiving the Data\n"
            "(i) Name of the Nodal Dept\n"
            "(ii) Date of Receipt of Request from Nodal Department.(Attach Copy of the letter)"
            , ""
        ),
        (
            "Details of Data Sharing\n"
            "(i) Date\n"
            "(ii) Details of the Officer receiving the data\n\n"
            "(iii) Mode of data sharing\n\n"
            "(iv) Signature of Officer receiving the data\n\n", ""
        ),
        ("Remarks", ""),
    ]

    for row, (label, value) in zip(table2.rows, table2_fields):
        row.cells[0].text = label
        row.cells[1].text = value or ""

    doc.add_paragraph("\n")


    # Create table
    sign_table = doc.add_table(rows=1, cols=2)

    tbl = sign_table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    tblPr.append(tblBorders)

    # Left cell → Date
    left_cell = sign_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = 0  # LEFT
    left_para.add_run("Date:\n\n")
    left_para.add_run("______________")

    # Right cell → Signature
    right_cell = sign_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = 2  # RIGHT
    right_para.add_run("Name & Signature of State NIC Co-ordinate\n\n")
    right_para.add_run("__________________________")


    filename = f"Audit_DataSharing_{rec.id}.docx"
    filepath = os.path.join(folder, filename)
    doc.save(filepath)

    return filename


def extract_filled_details(docx_path):

    doc = Document(docx_path)
    extracted = {}

    # Scan tables for key/value rows to match labels used in create_prefilled_doc
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = (row.cells[0].text or "").strip()
            value = (row.cells[1].text or "").strip()

            if "Period Up-to which the data was shared" in label:
                extracted["period_shared"] = value
            elif "Probable size of database after import" in label:
                extracted["db_size"] = value
            elif "Postgresql Version in which the Database backup was taken" in label:
                extracted["postgres_version"] = value
            elif label.strip().lower() == "coordinator":
                extracted["coordinator_name"] = value

    return extracted


def verify_pdf_signature(path):
    try:
        with fitz.open(path) as doc:
            for i in range(doc.page_count):
                page = doc.load_page(i)
                annots = page.annots()
                if not annots:
                    continue
                for a in annots:
                    # Heuristic: 20 used to represent signature annotations in some contexts
                    if hasattr(a, "type") and a.type and a.type[0] == 20:
                        return True
    except Exception:
        # On parsing failure, treat as unsigned
        return False
    return False


def verify_docx_signature(path):
    doc = Document(path)

    # Check for signature keywords
    signature_words = ["Signed", "Signature", "Authorized", "Seal"]

    for p in doc.paragraphs:
        if any(word.lower() in p.text.lower() for word in signature_words):
            return True

    # Check for images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            return True

    return False

def verify_signature(path):
    if path.lower().endswith(".pdf"):
        return verify_pdf_signature(path)
    if path.lower().endswith(".docx"):
        return verify_docx_signature(path)
    return False
