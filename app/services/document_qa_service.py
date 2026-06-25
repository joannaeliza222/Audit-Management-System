import os
import uuid
import hashlib
import magic
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
from flask import current_app
from cryptography.fernet import Fernet
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract
import io
import json
import logging
from typing import List, Dict, Optional, Tuple

from app import db
from ..models import User
from ..document_qa_models import (
    SecureDocument, QASession, QAConversation, 
    DocumentAccessLog, DocumentShare, DocumentStatus, QADocumentChunk
)
from ..utils.embeddings import get_embedding
from ..utils.security import generate_encryption_key, encrypt_data, decrypt_data


logger = logging.getLogger(__name__)


class DocumentQAService:
    """Service for secure document Q&A operations"""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt', 'md', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'}
    ALLOWED_MIME_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
        'text/plain',
        'text/markdown',
        'image/png',
        'image/jpeg',
        'image/tiff',
        'image/bmp'
    }
    MAX_FILE_SIZE_MB = 50
    CHUNK_SIZE = 1000  # Characters per chunk
    CHUNK_OVERLAP = 200  # Characters overlap between chunks
    
    def __init__(self):
        self.encryption_key = None
        self._init_encryption()
    
    def _init_encryption(self):
        """Initialize encryption for document storage"""
        try:
            key_str = current_app.config.get('DOCUMENT_ENCRYPTION_KEY')
            if key_str:
                self.encryption_key = key_str.encode()
            else:
                logger.warning("Document encryption key not configured")
        except Exception as e:
            logger.error(f"Failed to initialize encryption: {e}")
    
    def validate_file(self, file, filename: str) -> Tuple[bool, str]:
        """Validate uploaded file"""
        if not file or not filename:
            return False, "No file provided"
        
        # Check file extension
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        if ext not in self.ALLOWED_EXTENSIONS:
            return False, f"File type .{ext} not allowed"
        
        # Check file size
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        
        if size > self.MAX_FILE_SIZE_MB * 1024 * 1024:
            return False, f"File size exceeds {self.MAX_FILE_SIZE_MB}MB limit"
        
        # Check MIME type
        try:
            file_content = file.read(1024)
            file.seek(0)
            mime_type = magic.from_buffer(file_content, mime=True)
            
            if mime_type not in self.ALLOWED_MIME_TYPES:
                return False, f"MIME type {mime_type} not allowed"
        except Exception as e:
            logger.error(f"MIME type detection failed: {e}")
            return False, "File validation failed"
        
        return True, "File valid"
    
    def upload_document(self, user_id: int, file, filename: str) -> Tuple[Optional[SecureDocument], str]:
        """Upload and process a document"""
        # Validate file
        is_valid, error_msg = self.validate_file(file, filename)
        if not is_valid:
            return None, error_msg
        
        try:
            # Create document record
            document = SecureDocument(
                user_id=user_id,
                original_filename=filename,
                stored_filename=self._generate_secure_filename(filename),
                file_size=file.tell(),
                mime_type=magic.from_buffer(file.read(1024), mime=True),
                status=DocumentStatus.uploading
            )
            file.seek(0)
            
            # Calculate file hash
            file_content = file.read()
            document.file_hash = document.calculate_file_hash(file_content)
            file.seek(0)
            
            # Check for duplicate uploads
            existing = SecureDocument.query.filter_by(
                user_id=user_id,
                file_hash=document.file_hash,
                status=DocumentStatus.ready
            ).first()
            
            if existing:
                return None, "Document already exists"
            
            # Save file
            file_path = self._save_file(file, document.stored_filename)
            document.file_path = file_path
            
            db.session.add(document)
            db.session.commit()
            
            # Log upload
            self._log_access(document.id, user_id, 'upload')
            
            # Start processing in background
            self._process_document_async(document.id)
            
            return document, "Document uploaded successfully"
            
        except Exception as e:
            logger.error(f"Document upload failed: {e}")
            db.session.rollback()
            return None, f"Upload failed: {str(e)}"
    
    def _generate_secure_filename(self, original_filename: str) -> str:
        """Generate secure filename"""
        ext = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else ''
        return f"{uuid.uuid4().hex}.{ext}" if ext else uuid.uuid4().hex
    
    def _save_file(self, file, filename: str) -> str:
        """Save file to secure storage"""
        upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'documents')
        os.makedirs(upload_dir, exist_ok=True)
        
        file_path = os.path.join(upload_dir, filename)
        
        # Encrypt file if encryption is enabled
        if self.encryption_key:
            file_content = file.read()
            encrypted_content = encrypt_data(file_content, self.encryption_key)
            
            with open(file_path, 'wb') as f:
                f.write(encrypted_content)
        else:
            file.save(file_path)
        
        return file_path
    
    def _process_document_async(self, document_id: int):
        """Process document asynchronously"""
        try:
            document = SecureDocument.query.get(document_id)
            if not document:
                return
            
            document.status = DocumentStatus.processing
            document.processing_started_at = datetime.utcnow()
            db.session.commit()
            
            # Extract text
            text_content = self._extract_text(document)
            
            if not text_content:
                document.status = DocumentStatus.failed
                document.processing_error = "Failed to extract text"
                db.session.commit()
                return
            
            # Create chunks
            chunks = self._create_chunks(document, text_content)
            
            # Generate embeddings
            self._generate_embeddings(chunks)
            
            # Update document status
            document.status = DocumentStatus.ready
            document.processing_completed_at = datetime.utcnow()
            document.extracted_text_length = len(text_content)
            document.word_count = len(text_content.split())
            db.session.commit()
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            document = SecureDocument.query.get(document_id)
            if document:
                document.status = DocumentStatus.failed
                document.processing_error = str(e)
                db.session.commit()
    
    def _extract_text(self, document: SecureDocument) -> str:
        """Extract text from document"""
        try:
            # Decrypt file if encrypted
            if self.encryption_key:
                with open(document.file_path, 'rb') as f:
                    encrypted_content = f.read()
                file_content = decrypt_data(encrypted_content, self.encryption_key)
            else:
                with open(document.file_path, 'rb') as f:
                    file_content = f.read()
            
            # Extract based on MIME type
            if document.mime_type == 'application/pdf':
                return self._extract_pdf_text(file_content)
            elif 'wordprocessingml' in document.mime_type:
                return self._extract_docx_text(file_content)
            elif document.mime_type.startswith('text/'):
                return file_content.decode('utf-8', errors='ignore')
            elif document.mime_type.startswith('image/'):
                return self._extract_image_text(file_content)
            else:
                return ""
                
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return ""
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            doc = fitz.open(stream=file_content)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return ""
    
    def _extract_image_text(self, file_content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"OCR text extraction failed: {e}")
            return ""
    
    def _create_chunks(self, document: SecureDocument, text: str) -> List[QADocumentChunk]:
        """Create text chunks"""
        chunks = []
        
        if not text.strip():
            return chunks
        
        # Split into chunks with overlap
        text_length = len(text)
        start = 0
        chunk_index = 0
        
        while start < text_length:
            end = min(start + self.CHUNK_SIZE, text_length)
            
            # Try to break at sentence boundary
            if end < text_length:
                for i in range(end, max(start, end - 200), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                # Use QADocumentChunk model
                chunk = QADocumentChunk(
                    document_id=document.id,
                    user_id=document.user_id,
                    chunk_text=chunk_text,
                    chunk_index=chunk_index,
                    chunk_type='text'
                )
                chunks.append(chunk)
                chunk_index += 1
            
            start = end - self.CHUNK_OVERLAP if end < text_length else text_length
        
        # Save chunks
        for chunk in chunks:
            db.session.add(chunk)
        
        db.session.commit()
        return chunks
    
    def _generate_embeddings(self, chunks: List[QADocumentChunk]):
        """Generate embeddings for chunks"""
        try:
            texts = [chunk.chunk_text for chunk in chunks]
            embeddings = get_embedding(texts)
            
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding
            
            db.session.commit()
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise
    
    def create_qa_session(self, user_id: int, document_id: int, session_name: str = None) -> QASession:
        """Create Q&A session"""
        session = QASession(
            user_id=user_id,
            document_id=document_id,
            session_name=session_name
        )
        session.generate_session_token()
        
        db.session.add(session)
        db.session.commit()
        
        self._log_access(document_id, user_id, 'qa_session_created')
        return session
    
    def ask_question(self, session_id: int, question: str) -> Tuple[str, float, List[Dict]]:
        """Ask question about document"""
        try:
            session = QASession.query.get(session_id)
            if not session:
                return "Session not found", 0.0, []
            
            # Generate question embedding
            question_embedding = get_embedding([question])[0]
            
            # Find relevant chunks
            relevant_chunks = self._find_relevant_chunks(
                session.document_id, 
                session.user_id, 
                question_embedding
            )
            
            if not relevant_chunks:
                return "No relevant information found", 0.0, []
            
            # Generate answer using LLM
            answer, confidence = self._generate_answer(question, relevant_chunks)
            
            # Save conversation
            conversation = QAConversation(
                session_id=session_id,
                user_id=session.user_id,
                question=question,
                question_embedding=question_embedding,
                answer=answer,
                confidence_score=confidence,
                relevant_chunks=json.dumps([chunk.id for chunk in relevant_chunks]),
                context_length=sum(len(chunk.chunk_text) for chunk in relevant_chunks)
            )
            
            db.session.add(conversation)
            
            # Update session
            session.question_count += 1
            session.update_activity()
            
            db.session.commit()
            
            # Format sources
            sources = [
                {
                    'chunk_id': chunk.id,
                    'text': chunk.chunk_text[:200] + '...' if len(chunk.chunk_text) > 200 else chunk.chunk_text,
                    'page': chunk.page_number
                }
                for chunk in relevant_chunks
            ]
            
            self._log_access(session.document_id, session.user_id, 'question_asked')
            
            return answer, confidence, sources
            
        except Exception as e:
            logger.error(f"Question answering failed: {e}")
            return f"Error: {str(e)}", 0.0, []
    
    def _find_relevant_chunks(self, document_id: int, user_id: int, question_embedding, top_k: int = 5) -> List[QADocumentChunk]:
        """Find relevant chunks using vector similarity"""
        try:
            from sqlalchemy import select
            from pgvector.sqlalchemy import Vector
            
            # Vector similarity search
            query = select(QADocumentChunk).where(
                QADocumentChunk.document_id == document_id,
                QADocumentChunk.user_id == user_id
            ).order_by(
                QADocumentChunk.embedding.l2_distance(question_embedding)
            ).limit(top_k)
            
            result = db.session.execute(query).scalars().all()
            return list(result)
            
        except Exception as e:
            logger.error(f"Chunk search failed: {e}")
            return []
    
    def _generate_answer(self, question: str, chunks: List[QADocumentChunk]) -> Tuple[str, float]:
        """Generate answer using LLM"""
        try:
            # For now, use a simple approach - concatenate relevant chunks
            context = "\n\n".join([chunk.chunk_text for chunk in chunks])
            
            # This is where you would integrate with OpenAI or other LLM
            # For now, return a simple response
            answer = f"Based on the document, here's what I found:\n\n{context[:500]}..."
            confidence = 0.8
            
            return answer, confidence
            
        except Exception as e:
            logger.error(f"Answer generation failed: {e}")
            return "I'm sorry, I couldn't generate an answer.", 0.0
    
    def _log_access(self, document_id: int, user_id: int, action: str, additional_data: Dict = None):
        """Log document access"""
        try:
            log = DocumentAccessLog(
                document_id=document_id,
                user_id=user_id,
                action=action,
                additional_data=json.dumps(additional_data) if additional_data else None
            )
            db.session.add(log)
            db.session.commit()
        except Exception as e:
            logger.error(f"Access logging failed: {e}")
    
    def get_user_documents(self, user_id: int, status: DocumentStatus = None) -> List[SecureDocument]:
        """Get user's documents"""
        query = SecureDocument.query.filter_by(user_id=user_id)
        if status:
            query = query.filter_by(status=status)
        return query.order_by(SecureDocument.created_at.desc()).all()
    
    def delete_document(self, document_id: int, user_id: int) -> bool:
        """Delete document"""
        try:
            document = SecureDocument.query.filter_by(id=document_id, user_id=user_id).first()
            if not document:
                return False
            
            # Soft delete
            document.soft_delete()
            
            # Log deletion
            self._log_access(document_id, user_id, 'delete')
            
            db.session.commit()
            return True
            
        except Exception as e:
            logger.error(f"Document deletion failed: {e}")
            db.session.rollback()
            return False
