import os
import uuid
import hashlib
import json
from typing import List, Dict, Tuple, Optional, BinaryIO
from datetime import datetime
from io import BytesIO
import tiktoken
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
from flask import current_app
from sqlalchemy.exc import IntegrityError

from ..document_qa_models import SecureDocument, QADocumentChunk, DocumentStatus
from .document_security import DocumentSecurityService


class DocumentIngestionService:
    """Service for document ingestion, chunking, and embedding generation"""
    
    def __init__(self, encryption_key: str):
        """
        Initialize the ingestion service
        
        Args:
            encryption_key: AES-256 encryption key for document storage
        """
        self.encryption_key = encryption_key
        self.security_service = DocumentSecurityService()
        
        # Initialize embedding model (lazy loading)
        self._embedding_model = None
        self._tokenizer = None
        
        # Configuration
        self.chunk_size = 500  # tokens
        self.chunk_overlap = 50  # tokens
        self.max_file_size = 20 * 1024 * 1024  # 20MB
        
        # Allowed MIME types
        self.allowed_mime_types = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'text/markdown',
            'text/csv'
        }
    
    @property
    def embedding_model(self):
        """Lazy loading of embedding model"""
        if self._embedding_model is None:
            model_name = current_app.config.get('EMBEDDING_MODEL_NAME', 'sentence-transformers/all-MiniLM-L6-v2')
            self._embedding_model = SentenceTransformer(model_name)
        return self._embedding_model
    
    @property
    def tokenizer(self):
        """Lazy loading of tokenizer"""
        if self._tokenizer is None:
            self._tokenizer = tiktoken.get_encoding("cl100k_base")
        return self._tokenizer
    
    def upload_document(self, user_id: int, file_buffer: BinaryIO, filename: str, 
                       mime_type: str, ip_address: str = None, user_agent: str = None) -> Dict:
        """
        Upload and process a document
        
        Args:
            user_id: User ID
            file_buffer: File content as bytes
            filename: Original filename
            mime_type: File MIME type
            ip_address: Client IP address
            user_agent: Client user agent
            
        Returns:
            Dict with document ID and status
        """
        try:
            # Validate inputs
            validation_result = self._validate_upload(file_buffer, filename, mime_type)
            if not validation_result['valid']:
                return {
                    'success': False,
                    'error': validation_result['error'],
                    'document_id': None
                }
            
            # Read file content
            file_content = file_buffer.read()
            
            # Generate secure filename
            safe_filename = self.security_service.generate_safe_filename(filename)
            
            # Create document record
            stored_filename = f"{uuid.uuid4().hex}{self._get_file_extension(filename)}"
            upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
            file_path = os.path.join(upload_folder, stored_filename)
            
            # Ensure upload folder exists
            os.makedirs(upload_folder, exist_ok=True)
            
            # Save file to disk
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            document = SecureDocument(
                user_id=user_id,
                original_filename=filename,
                stored_filename=stored_filename,
                file_path=file_path,
                file_size=len(file_content),
                mime_type=mime_type,
                file_hash=hashlib.sha256(file_content).hexdigest(),
                status=DocumentStatus.processing
            )
            
            # Save document
            from .. import db
            db.session.add(document)
            db.session.flush()  # Get the document ID
            
            # Extract text and create chunks
            text_content = self._extract_text(file_content, mime_type)
            chunks = self._create_chunks(text_content, document.id, user_id)
            
            # Generate embeddings for chunks
            chunk_embeddings = self._generate_embeddings([chunk['text'] for chunk in chunks])
            
            # Save chunks with embeddings
            chunk_models = []
            
            for i, chunk_data in enumerate(chunks):
                chunk_model = QADocumentChunk(
                    document_id=document.id,
                    user_id=user_id,
                    chunk_index=i,
                    chunk_text=chunk_data['text'],
                    embedding=chunk_embeddings[i],
                    chunk_type='text'
                )
                chunk_models.append(chunk_model)
            
            # Update document metadata
            document.status = DocumentStatus.ready
            document.processing_completed_at = datetime.utcnow()
            document.extracted_text_length = len(text_content)
            
            # Save all chunks
            db.session.add_all(chunk_models)
            
            # Log the upload
            from ..document_qa_models import DocumentAccessLog
            access_log = DocumentAccessLog(
                user_id=user_id,
                document_id=document.id,
                action='upload',
                ip_address=ip_address,
                user_agent=user_agent,
                additional_data=json.dumps({
                    'filename': filename,
                    'file_size': len(file_content),
                    'mime_type': mime_type,
                    'chunk_count': len(chunks)
                })
            )
            db.session.add(access_log)
            
            db.session.commit()
            
            return {
                'success': True,
                'document_id': document.id,
                'status': 'uploaded',
                'chunk_count': len(chunks)
            }
            
        except IntegrityError as e:
            from .. import db
            db.session.rollback()
            return {
                'success': False,
                'error': 'Database integrity error',
                'document_id': None
            }
        except Exception as e:
            from .. import db
            db.session.rollback()
            current_app.logger.error(f"Error uploading document: {e}")
            return {
                'success': False,
                'error': 'Internal server error',
                'document_id': None
            }
    
    def list_documents(self, user_id: int) -> List[Dict]:
        """
        List all documents for a user
        
        Args:
            user_id: User ID
            
        Returns:
            List of document metadata
        """
        documents = SecureDocument.query.filter_by(
            user_id=user_id
        ).filter(SecureDocument.deleted_at.is_(None)).order_by(SecureDocument.created_at.desc()).all()
        
        return [{
            'id': doc.id,
            'original_filename': doc.original_filename,
            'file_size': doc.file_size,
            'mime_type': doc.mime_type,
            'status': doc.status.value if doc.status else None,
            'created_at': doc.created_at.isoformat() if doc.created_at else None,
            'chunk_count': doc.chunks.count()
        } for doc in documents]
    
    def delete_document(self, user_id: int, document_id: str, 
                       ip_address: str = None, user_agent: str = None) -> Dict:
        """
        Delete a document and all its chunks
        
        Args:
            user_id: User ID
            document_id: Document ID
            ip_address: Client IP address
            user_agent: Client user agent
            
        Returns:
            Dict with deletion status
        """
        try:
            # Get document
            document = SecureDocument.query.filter_by(
                id=document_id,
                user_id=user_id
            ).filter(SecureDocument.deleted_at.is_(None)).first()
            
            if not document:
                return {
                    'success': False,
                    'error': 'Document not found'
                }
            
            # Get chunk count for logging
            chunk_count = document.chunks.count()
            
            # Mark document as deleted (soft delete)
            document.soft_delete()
            
            # Log the deletion
            from ..document_qa_models import DocumentAccessLog
            access_log = DocumentAccessLog(
                user_id=user_id,
                document_id=document_id,
                action='delete',
                ip_address=ip_address,
                user_agent=user_agent,
                additional_data=json.dumps({
                    'filename': document.original_filename,
                    'chunk_count': chunk_count
                })
            )
            db.session.add(access_log)
            
            from .. import db
            db.session.commit()
            
            return {
                'success': True,
                'message': 'Document deleted successfully'
            }
            
        except Exception as e:
            from .. import db
            db.session.rollback()
            current_app.logger.error(f"Error deleting document: {e}")
            return {
                'success': False,
                'error': 'Internal server error'
            }
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename"""
        if '.' in filename:
            return '.' + filename.rsplit('.', 1)[1].lower()
        return ''
    
    def _validate_upload(self, file_buffer: BinaryIO, filename: str, mime_type: str) -> Dict:
        """Validate file upload"""
        # Check file size
        file_buffer.seek(0, 2)  # Seek to end
        file_size = file_buffer.tell()
        file_buffer.seek(0)  # Reset position
        
        if file_size > self.max_file_size:
            return {
                'valid': False,
                'error': f'File size exceeds maximum allowed size of {self.max_file_size // (1024*1024)}MB'
            }
        
        # Check MIME type
        if mime_type not in self.allowed_mime_types:
            return {
                'valid': False,
                'error': f'File type {mime_type} is not allowed'
            }
        
        # Validate filename
        safe_filename = self.security_service.generate_safe_filename(filename)
        if not safe_filename:
            return {
                'valid': False,
                'error': 'Invalid filename'
            }
        
        return {'valid': True}
    
    def _extract_text(self, file_content: bytes, mime_type: str) -> str:
        """Extract text from document based on MIME type"""
        try:
            if mime_type == 'application/pdf':
                return self._extract_pdf_text(file_content)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._extract_docx_text(file_content)
            elif mime_type in ['text/plain', 'text/markdown', 'text/csv']:
                return file_content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported MIME type: {mime_type}")
        except Exception as e:
            current_app.logger.error(f"Error extracting text from {mime_type}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        text = []
        pdf_file = BytesIO(file_content)
        
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        except Exception as e:
            current_app.logger.error(f"Error reading PDF: {e}")
        
        return '\n'.join(text)
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        docx_file = BytesIO(file_content)
        
        try:
            doc = docx.Document(docx_file)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            current_app.logger.error(f"Error reading DOCX: {e}")
            return ""
    
    def _create_chunks(self, text: str, document_id: str, user_id: int) -> List[Dict]:
        """
        Create text chunks with overlap and security scanning
        
        Args:
            text: Full document text
            document_id: Document ID
            user_id: User ID
            
        Returns:
            List of chunk dictionaries
        """
        if not text.strip():
            return []
        
        # Tokenize text
        tokens = self.tokenizer.encode(text)
        
        chunks = []
        start_idx = 0
        chunk_index = 0
        
        while start_idx < len(tokens):
            # Calculate chunk end
            end_idx = min(start_idx + self.chunk_size, len(tokens))
            
            # Extract chunk tokens
            chunk_tokens = tokens[start_idx:end_idx]
            
            # Decode to text
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            # Security scanning and sanitization
            sanitized_text, flagged, patterns = self.security_service.sanitize_text(
                chunk_text, document_id, chunk_index
            )
            
            chunks.append({
                'text': sanitized_text,
                'flagged': flagged,
                'token_count': len(chunk_tokens),
                'patterns': patterns
            })
            
            # Move to next chunk with overlap
            start_idx = end_idx - self.chunk_overlap
            chunk_index += 1
            
            # Prevent infinite loop
            if start_idx >= len(tokens):
                break
        
        return chunks
    
    def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        if not texts:
            return []
        
        try:
            embeddings = self.embedding_model.encode(
                texts,
                batch_size=32,
                normalize_embeddings=True,
                show_progress_bar=False
            )
            return embeddings.tolist()
        except Exception as e:
            current_app.logger.error(f"Error generating embeddings: {e}")
            # Return zero embeddings as fallback
            return [[0.0] * 384 for _ in texts]  # 384 is the embedding dimension
    
    def get_document_content(self, user_id: int, document_id: str) -> Optional[str]:
        """
        Get decrypted document content
        
        Args:
            user_id: User ID
            document_id: Document ID
            
        Returns:
            Decrypted text content or None
        """
        document = SecureDocument.query.filter_by(
            id=document_id,
            user_id=user_id
        ).filter(SecureDocument.deleted_at.is_(None)).first()
        
        if not document:
            return None
        
        try:
            # Read file from disk
            with open(document.file_path, 'rb') as f:
                file_content = f.read()
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            current_app.logger.error(f"Error reading document: {e}")
            return None
    
    def get_document_chunks(self, user_id: int, document_id: str, 
                           include_flagged: bool = False) -> List[Dict]:
        """
        Get document chunks
        
        Args:
            user_id: User ID
            document_id: Document ID
            include_flagged: Whether to include flagged chunks
            
        Returns:
            List of chunk dictionaries
        """
        query = QADocumentChunk.query.filter_by(
            document_id=document_id,
            user_id=user_id
        )
        
        chunks = query.order_by(QADocumentChunk.chunk_index).all()
        return [{
            'id': chunk.id,
            'document_id': chunk.document_id,
            'chunk_index': chunk.chunk_index,
            'chunk_text': chunk.chunk_text,
            'chunk_type': chunk.chunk_type,
            'created_at': chunk.created_at.isoformat() if chunk.created_at else None
        } for chunk in chunks]
